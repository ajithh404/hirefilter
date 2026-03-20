"""
file_parser.py — Binary file text extraction (PDF, DOCX)
Author: written by hand

Keeps file-format concerns completely separate from the text-processing
logic in parser.py. This module only does one thing: take bytes + a
filename, return plain text. The rest of the pipeline doesn't care
what format the file was.

Supported formats:
  .pdf  — PyMuPDF (fitz): fast, no external poppler dependency
  .docx — python-docx: reads paragraphs + tables
  .txt  — decoded as UTF-8 with fallback to latin-1
  .md   — same as .txt

Why PyMuPDF over pdfplumber / PyPDF2?
  - No subprocess calls, pure Python binding to MuPDF
  - Better at handling multi-column layouts and ligatures
  - Maintains reading order better than most alternatives
"""

import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


# ── PDF extraction ─────────────────────────────────────────────────────────────

def extract_pdf(data: bytes) -> str:
    """
    Extract plain text from a PDF byte string using PyMuPDF.
    Joins pages with a double newline to preserve rough structure.
    """
    try:
        import fitz  # PyMuPDF — imported here so the module loads even if not installed
    except ImportError:
        raise RuntimeError(
            "PyMuPDF is not installed. Run: pip install pymupdf"
        )

    pages = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            # get_text("text") preserves reading order, strips annotations
            text = page.get_text("text").strip()
            if text:
                pages.append(text)

    if not pages:
        logger.warning("PDF had no extractable text — may be a scanned image")
        return ""

    return "\n\n".join(pages)


# ── DOCX extraction ────────────────────────────────────────────────────────────

def extract_docx(data: bytes) -> str:
    """
    Extract text from a .docx byte string using python-docx.
    Reads both body paragraphs and table cells so skill grids aren't skipped.
    """
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        )

    doc = Document(io.BytesIO(data))
    lines = []

    # Body paragraphs (preserves headings, bullet points, etc.)
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # Table cells — many resumes put skills in a table grid
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                lines.append(" | ".join(row_text))

    return "\n".join(lines)


# ── Plain text ─────────────────────────────────────────────────────────────────

def extract_text(data: bytes) -> str:
    """
    Decode bytes as UTF-8, falling back to latin-1 for legacy files.
    Handles .txt and .md files.
    """
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("latin-1", errors="replace")


# ── Public dispatcher ──────────────────────────────────────────────────────────

def extract_file_text(filename: str, data: bytes) -> str:
    """
    Route a file to the correct extractor based on its extension.
    Returns plain text, or raises ValueError for unsupported formats.

    Args:
        filename: original filename (used for extension detection only)
        data:     raw file bytes

    Returns:
        Extracted plain text string
    """
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        return extract_pdf(data)
    elif suffix == ".docx":
        return extract_docx(data)
    elif suffix in (".txt", ".md", ".text"):
        return extract_text(data)
    else:
        raise ValueError(
            f"Unsupported file type: '{suffix}'. "
            "Accepted formats: .pdf, .docx, .txt, .md"
        )
